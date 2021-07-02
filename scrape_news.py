import requests, sys, wget
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

url = sys.argv[1]

#Check if argv[1] is a URL.

#Get the article
r = requests.get(url)
#Parse the article
soup = BeautifulSoup(r.content,"lxml")

doc = Document()
doc.add_heading(soup.title.string[:-11], 0)

picture = soup.find("img").get_attribute_list("src")[0]
filename = wget.download(picture)

doc.add_picture(filename,width=Inches(6.0))
for p in soup.article.find_all("p"): 
    if p.string == None:
        continue
    elif p.find("b") != None:
        paragraph = doc.add_paragraph()
        paragraph.add_run(p.find("b").string).bold = True
        continue
    doc.add_paragraph(p.string)
      

    
doc.save('demo.docx')




